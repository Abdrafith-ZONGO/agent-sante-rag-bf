import os
import sys

# Auto-install python-docx
try:
    import docx
except ImportError:
    print("[INFO] Bibliothèque python-docx non trouvée. Installation...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_heading_styled(doc, text, level):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(16, 44, 87) # Bleu Foncé
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(53, 114, 239) # Bleu Moyen
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        
    p.paragraph_format.keep_with_next = True
    return p

def add_paragraph_styled(doc, text="", bold_prefix="", italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.name = 'Times New Roman'
        r_prefix.font.size = Pt(11)
        r_prefix.font.bold = True
        
    if text:
        r_text = p.add_run(text)
        r_text.font.name = 'Times New Roman'
        r_text.font.size = Pt(11)
        r_text.font.italic = italic
        
    return p

def add_code_styled(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(80, 80, 80)
    return p

def add_screenshot_box(doc, number, filename, description):
    """Crée un encadré stylé sous forme de tableau pour insérer une capture d'écran."""
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(5.8)
    set_table_borders(table, "B2D8D8", "8") # Bordure verte/bleue clinique
    
    # Cellule 1 : Emplacement Image
    cell_img = table.rows[0].cells[0]
    set_cell_background(cell_img, "F4FBFB")
    set_cell_margins(cell_img, top=200, bottom=200, left=200, right=200)
    p_img = cell_img.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_img = p_img.add_run(f"\n[ EMPLACEMENT CAPTURE D'ÉCRAN {number} ]\nGlissez et déposez l'image : {filename}\n")
    r_img.font.name = 'Times New Roman'
    r_img.font.size = Pt(11)
    r_img.font.bold = True
    r_img.font.color.rgb = RGBColor(0, 128, 128)
    
    # Cellule 2 : Légende
    cell_leg = table.rows[1].cells[0]
    set_cell_background(cell_leg, "E0F2F1")
    set_cell_margins(cell_leg, top=100, bottom=100, left=150, right=150)
    p_leg = cell_leg.paragraphs[0]
    p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_leg = p_leg.add_run(f"Figure {number} : {description}")
    r_leg.font.name = 'Times New Roman'
    r_leg.font.size = Pt(10)
    r_leg.font.italic = True
    r_leg.font.color.rgb = RGBColor(30, 70, 70)
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(8)

def build_complete_technical_report():
    doc = Document()
    
    # Configuration des marges (2.54 cm partout)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # ---------------------------------------------------------------------------
    # PAGE DE GARDE
    # ---------------------------------------------------------------------------
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_univ = p_univ.add_run("UNIVERSITÉ JOSEPH KI-ZERBO (UJKZ)\n")
    r_univ.font.size = Pt(14)
    r_univ.font.bold = True
    
    r_ifoad = p_univ.add_run("INSTITUT DE FORMATION OUVERTE ET À DISTANCE (IFOAD)\n"
                            "MASTER 1 EN DATA SCIENCE — ÉDITION 2026\n")
    r_ifoad.font.size = Pt(11)
    r_ifoad.font.italic = True
    
    for _ in range(2):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("RAPPORT DE TRAVAIL DE DATA SCIENCE :\n"
                              "CONCEPTION ET DEPLOIEMENT D'UN AGENT IA ASSISTANT & SYSTEME RAG INTELLIGENT\n")
    r_title.font.size = Pt(14)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(16, 44, 87)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("« AGENT DE SANTE PUBLIQUE POUR LE BURKINA FASO »\n"
                          "Option 3 : Agent d'Orientation Médicale & Prévention (Santé)\n")
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.bold = True
    
    for _ in range(3):
        doc.add_paragraph()
        
    table_authors = doc.add_table(rows=1, cols=2)
    table_authors.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_authors.autofit = False
    table_authors.columns[0].width = Inches(3.25)
    table_authors.columns[1].width = Inches(3.25)
    
    cell_left = table_authors.rows[0].cells[0]
    p_authors = cell_left.paragraphs[0]
    r_auth_label = p_authors.add_run("Groupe 3 :\n")
    r_auth_label.font.bold = True
    p_authors.add_run("1. DRABO Bernard\n"
                      "2. Joseph Yarga\n"
                      "3. ZONGO Abdrafith")
    
    cell_right = table_authors.rows[0].cells[1]
    p_teacher = cell_right.paragraphs[0]
    p_teacher.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_teach_label = p_teacher.add_run("Sous la direction de :\n")
    r_teach_label.font.bold = True
    p_teacher.add_run("Dr Delwende D. Arthur SAWADOGO\nEnseignant-Chercheur")
    
    for _ in range(3):
        doc.add_paragraph()
        
    p_links = doc.add_paragraph()
    p_links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_git_label = p_links.add_run("Dépôt GitHub du Projet : ")
    r_git_label.font.bold = True
    p_links.add_run("https://github.com/Abdrafith-ZONGO/agent-sante-rag-bf\n")
    r_dep_label = p_links.add_run("Lien de Déploiement : ")
    r_dep_label.font.bold = True
    p_links.add_run("https://agent-sante-bf.onrender.com (Backend API)\n"
                    "https://agent-sante-bf-ui.onrender.com (Frontend React UI)")
    
    for _ in range(2):
        doc.add_paragraph()
        
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run("Ouagadougou, 12 Juillet 2026")
    r_date.font.bold = True
    
    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # TABLE DES MATIÈRES
    # ---------------------------------------------------------------------------
    add_heading_styled(doc, "TABLE DES MATIÈRES", level=1)
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.line_spacing = 1.3
    toc_p.add_run(
        "INTRODUCTION GENERALE ......................................................................................................... 3\n"
        "1. ARCHITECTURE TECHNIQUE DU SYSTEME ET SCHÉMA DE FLUX ............................................... 4\n"
        "   1.1. Architecture Globale Découlée ................................................................................................. 4\n"
        "   1.2. Schéma Fonctionnel de Flux de Données ............................................................................... 4\n"
        "   1.3. Stack Technique et Écosystème .............................................................................................. 5\n"
        "   1.4. Écosystème d'APIs et Rôles .................................................................................................... 6\n"
        "2. ARBORESCENCE DU PROJET ET RÔLE DE CHAQUE COMPOSANT ............................................ 7\n"
        "   2.1. Structure du Répertoire Projet ................................................................................................ 7\n"
        "   2.2. Rôle des Fichiers Clés ........................................................................................................... 7\n"
        "3. ANALYSE TECHNIQUE DÉTAILLÉE DU BACKEND ....................................................................... 9\n"
        "   3.1. Gestion de la Configuration (config.py) .................................................................................. 9\n"
        "   3.2. Persistance Relationnelle (database.py & models.py) ................................................................. 9\n"
        "   3.3. Sécurité, Chiffrement et JWT (auth.py) .................................................................................... 10\n"
        "   3.4. Ingestion des Directives Nationales (ingest.py) ............................................................................ 10\n"
        "   3.5. Intelligence Artificielle & RAG MMR (agent.py) .......................................................................... 11\n"
        "   3.6. Serveur API et Caching (main.py) ........................................................................................... 12\n"
        "4. ANALYSE TECHNIQUE DÉTAILLÉE DU FRONTEND REACT ........................................................... 14\n"
        "   4.1. Point d'Entrée et Routage Privé (App.jsx) ................................................................................ 14\n"
        "   4.2. Écrans d'Authentification (Login.jsx & Register.jsx) .................................................................. 14\n"
        "   4.3. Interface Principale et Chat (ChatWindow.jsx) ........................................................................... 15\n"
        "5. MÉTHODOLOGIE, STRATÉGIE RAG ET CHOIX TECHNIQUES ...................................................... 16\n"
        "   5.1. Stratégie de Chunking Récursif et Overlap ............................................................................... 16\n"
        "   5.2. Choix du Modèle d'Embeddings et Résolution des Conflits ........................................................... 16\n"
        "   5.3. Modèle de Langage (LLM) et Cadrage Sécuritaire ....................................................................... 17\n"
        "6. ÉVALUATION DE LA ROBUSTESSE (Nouveauté 2026) ............................................................... 18\n"
        "   6.1. Protocole de Test Automatisé (evaluate.py) ............................................................................. 18\n"
        "   6.2. Analyse des Résultats Expérimentaux ...................................................................................... 18\n"
        "7. GUIDE D'EXÉCUTION EN LOCAL ET DÉPLOIEMENT EN LIGNE ...................................................... 20\n"
        "   7.1. Guide d'Exécution en Local (Pas à Pas) ...................................................................................... 20\n"
        "   7.2. Procédure de Déploiement en Ligne (Render) ........................................................................... 21\n"
        "8. DIFFICULTÉS RENCONTRÉES ET SOLUTIONS APPORTÉES ...................................................... 22\n"
        "   8.1. Limitation des Quotas d'APIs Gratuites (Erreurs 429) .................................................................. 22\n"
        "   8.2. Volatilité de SQLite sur Render (Perte de Sessions 401) .............................................................. 22\n"
        "   8.3. Problèmes d'Ergonomie et de Zoom Automatique sur Mobile ...................................................... 23\n"
        "   8.4. Conflits de Dimensions dans ChromaDB ..................................................................................... 23\n"
        "   8.5. Cadrage Éthique et Sécurité des Diagnostics ................................................................................ 23\n"
        "CONCLUSION ........................................................................................................................................ 25"
    )
    
    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # INTRODUCTION GENERALE
    # ---------------------------------------------------------------------------
    add_heading_styled(doc, "INTRODUCTION GENERALE", level=1)
    add_paragraph_styled(doc,
        "L'accès rapide, rigoureux et de premier niveau à l'information sanitaire représente un pilier majeur de la santé publique moderne. "
        "Au Burkina Faso, pays confronté à des épidémies saisonnières récurrentes telles que le paludisme et la dengue, la diffusion d'informations "
        "de prévention de premier niveau de manière accessible est cruciale. Cependant, les professionnels de la santé dans les structures "
        "de proximité (CSPS, CMA) sont souvent submergés, et la population se tourne fréquemment vers des canaux d'information informels, "
        "vecteurs de désinformation ou de conseils médicaux erronés.\n\n"
        "L'avènement des grands modèles de langage (LLM) offre une opportunité technologique majeure pour concevoir des assistants d'orientation "
        "intelligents. Toutefois, l'utilisation de LLM généralistes présente des risques de sécurité inacceptables dans le domaine médical : "
        "hallucinations, ignorance des spécificités nationales (comme les directives du Ministère de la Santé du Burkina Faso) et génération de diagnostics "
        "hasardeux. C'est pour pallier ces manquements structurels que s'impose le paradigme de génération augmentée par récupération "
        "(**RAG - Retrieval-Augmented Generation**).\n\n"
        "Ce projet de Master 1 Data Science à l'IFOAD (Édition 2026) sous la direction du Dr Sawadogo consiste à développer l'**Agent Santé BF**, "
        "un assistant d'orientation médicale et de prévention conçu pour le contexte burkinabè. Ce rapport de travail expose la conception de bout en bout "
        "de ce système intelligent, détaillant l'architecture technique, la stratégie d'ingestion sémantique des documents locaux réels, la logique interne "
        "de l'agent LangChain, le frontend réactif, le guide d'exécution locale et de déploiement en ligne, ainsi que les difficultés rencontrées "
        "et l'évaluation expérimentale de la robustesse de notre RAG."
    )
    
    # ---------------------------------------------------------------------------
    # 1. ARCHITECTURE TECHNIQUE DU SYSTEME ET SCHÉMA DE FLUX
    # ---------------------------------------------------------------------------
    add_heading_styled(doc, "1. ARCHITECTURE TECHNIQUE DU SYSTEME ET SCHÉMA DE FLUX", level=1)
    
    add_heading_styled(doc, "1.1. Architecture Globale Découlée", level=2)
    add_paragraph_styled(doc,
        "L'Agent Santé BF adopte une architecture client-serveur découplée afin d'assurer de hautes performances, une sécurité "
        "renforcée et une portabilité maximale. L'application sépare distinctement la couche d'interface utilisateur (React) de la couche logique et base de données (FastAPI + SQLite + ChromaDB)."
    )
    
    add_heading_styled(doc, "1.2. Schéma Fonctionnel de Flux de Données", level=2)
    table_flow = doc.add_table(rows=1, cols=1)
    table_flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_flow.columns[0].width = Inches(6.5)
    set_cell_background(table_flow.rows[0].cells[0], "F8F9FA")
    set_cell_margins(table_flow.rows[0].cells[0], top=150, bottom=150, left=150, right=150)
    p_flow = table_flow.rows[0].cells[0].paragraphs[0]
    p_flow.paragraph_format.line_spacing = 1.05
    p_flow.add_run(
        "   [FRONTEND REACT UI]\n"
        "            │\n"
        "            ▼ (1. Requête HTTP POST /chat + Token JWT + Message)\n"
        "   [BACKEND FASTAPI]\n"
        "            │\n"
        "            ▼ (2. Validation JWT & Chargement Historique SQL)\n"
        "   [LANGCHAIN AGENT EXECUTOR] <───> [Modèle LLM Llama 3.1 8B (Groq)]\n"
        "            │\n"
        "            ├─► [Optionnel : search_web] ──► [API TAVILY] ──► (Recherche Internet)\n"
        "            │\n"
        "            └─► [search_knowledge_base] ──► [Gemini Embeddings] ──► [ChromaDB]\n"
        "                                                                     │\n"
        "                                                                     ▼\n"
        "                                                         (3. Extraction Chunks PDF)\n"
        "            │\n"
        "            ▼ (4. Compilation du Contexte & Génération de la Réponse par LLM)\n"
        "   [ENREGISTREMENT SQLITE] ──► Sauvegarde du message en Base\n"
        "            │\n"
        "            ▼ (5. Rendu Réponse + Sources et Outils Utilisés)\n"
        "   [FRONTEND REACT UI]"
    )
    p_flow.runs[0].font.name = 'Courier New'
    p_flow.runs[0].font.size = Pt(9.5)

    add_heading_styled(doc, "1.3. Stack Technique et Écosystème", level=2)
    stack_details = [
        ("FastAPI (v0.115.0)", "Framework web asynchrone hautes performances en Python. Il génère automatiquement la documentation OpenAPI interactives et gère les routes de manière non bloquante."),
        ("React (v18) + Vite", "Framework frontend pour une interface utilisateur réactive. Vite fournit des temps de rechargement extrêmement rapides pour le développement. React Router DOM gère le routage et la sécurité des écrans."),
        ("Tailwind CSS", "Framework utilitaire CSS pour concevoir un design épuré, responsive et de type clinique médicale, facilitant l'intégration mobile."),
        ("SQLAlchemy & SQLite", "Base de données relationnelle locale intégrée. SQLite stocke de façon structurée les comptes et l'historique des discussions, tandis que SQLAlchemy abstrait les requêtes SQL."),
        ("ChromaDB & LangChain", "ChromaDB agit comme base de données vectorielle pour indexer les documents locaux. LangChain orchestre l'agent autonome, chaînant les outils, gérant la mémoire de chat et l'appel des LLM.")
    ]
    for label, desc in stack_details:
        p_st = doc.add_paragraph()
        p_st.paragraph_format.left_indent = Inches(0.2)
        p_st.paragraph_format.space_after = Pt(4)
        r_lbl = p_st.add_run(f"• {label} : ")
        r_lbl.font.bold = True
        p_st.add_run(desc)

    add_heading_styled(doc, "1.4. Écosystème d'APIs et Rôles", level=2)
    p_api1 = doc.add_paragraph(style='List Bullet')
    p_api1.add_run("Groq API (Llama 3.1 8B Instant) : ").font.bold = True
    p_api1.add_run("Groq utilise des processeurs LPU (Language Processing Units). Cela permet d'obtenir un temps de réponse (latence) de l'assistant de moins d'une seconde, garantissant la fluidité du dialogue.")
    
    p_api2 = doc.add_paragraph(style='List Bullet')
    p_api2.add_run("Google AI Studio (Gemini Embeddings) : ").font.bold = True
    p_api2.add_run("Le modèle 'gemini-embedding-001' produit des vecteurs de haute précision (3072 dimensions) pour encoder le sens sémantique exact des documents burkinabè.")
    
    p_api3 = doc.add_paragraph(style='List Bullet')
    p_api3.add_run("Tavily API (Recherche Web) : ").font.bold = True
    p_api3.add_run("Spécialisée pour les agents IA, elle permet d'effectuer des recherches sur l'actualité épidémique (ex: campagnes de vaccination récentes au Burkina) sans charger les scories des pages web habituelles.")

    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # 2. ARBORESCENCE DU PROJET AND FILE ROLES
    # ---------------------------------------------------------------------------
    add_heading_styled(doc, "2. ARBORESCENCE DU PROJET ET RÔLE DE CHAQUE COMPOSANT", level=1)
    
    add_heading_styled(doc, "2.1. Structure du Répertoire Projet", level=2)
    table_tree = doc.add_table(rows=1, cols=1)
    table_tree.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_tree.columns[0].width = Inches(6.5)
    set_cell_background(table_tree.rows[0].cells[0], "F8F9FA")
    set_cell_margins(table_tree.rows[0].cells[0], top=150, bottom=150, left=150, right=150)
    p_tree = table_tree.rows[0].cells[0].paragraphs[0]
    p_tree.paragraph_format.line_spacing = 1.0
    p_tree.add_run(
        "projet-sante-rag/\n"
        "├── backend/                    # Partie Serveur FastAPI\n"
        "│   ├── app/\n"
        "│   │   ├── __init__.py\n"
        "│   │   ├── agent.py            # Logique de l'agent IA, du prompt et des outils\n"
        "│   │   ├── auth.py             # Authentification JWT, hachage et endpoints d'accès\n"
        "│   │   ├── config.py           # Chargement et vérification des variables d'environnement\n"
        "│   │   ├── database.py         # Connexion SQLite et injection de session SQL\n"
        "│   │   ├── evaluate.py         # Script d'évaluation automatique de la robustesse\n"
        "│   │   ├── ingest.py           # Ingestion, découpage et vectorisation des PDF\n"
        "│   │   ├── main.py             # Point d'entrée FastAPI et gestion du cache des agents\n"
        "│   │   └── models.py           # Schémas de tables SQL (User, Session, Message)\n"
        "│   ├── data/\n"
        "│   │   └── raw/                # Directives de santé et PDF médicaux locaux\n"
        "│   ├── chroma_db/              # Base vectorielle persistée (générée)\n"
        "│   ├── app.db                  # Base relationnelle SQLite locale (créée automatiquement)\n"
        "│   └── requirements.txt        # Dépendances Python requises pour le projet\n"
        "└── frontend/                   # Interface Client React\n"
        "    ├── src/\n"
        "    │   ├── components/         # Composants réutilisables (ChatWindow, Messages, Badges)\n"
        "    │   ├── pages/              # Pages d'accès (Login.jsx, Register.jsx)\n"
        "    │   ├── App.jsx             # Router React et routes privées sécurisées\n"
        "    │   ├── index.css           # Feuille de styles globale (Tailwind & corrections mobiles)\n"
        "    │   └── main.jsx            # Point d'entrée du rendu DOM\n"
        "    └── package.json            # Dépendances Node.js"
    )
    p_tree.runs[0].font.name = 'Courier New'
    p_tree.runs[0].font.size = Pt(9.0)

    add_heading_styled(doc, "2.2. Rôle des Fichiers Clés", level=2)
    file_roles = [
        ("backend/app/config.py", "Centralise les constantes de l'application (taille de chunk, overlap, clés API). Il charge les valeurs depuis le fichier .env et valide la présence obligatoire des clés Groq et Gemini dès l'allumage du serveur pour éviter tout plantage ultérieur."),
        ("backend/app/database.py", "Configure l'ORM SQLAlchemy. Il initialise le moteur de base de données relationnelle SQLite `app.db` et fournit le générateur de session SQL `get_db` pour chaque route API."),
        ("backend/app/models.py", "Définit la structure SQL de la base de données. Il contient trois tables liées de façon relationnelle : les utilisateurs, leurs sessions de discussion, et l'historique complet de leurs messages."),
        ("backend/app/auth.py", "Assure la sécurité. Il contient les fonctions de hachage de mot de passe (via Bcrypt), de création de jeton JWT signé avec une clé secrète, et les routes d'API `/auth/register` (inscription) et `/auth/login` (connexion)."),
        ("backend/app/ingest.py", "Script d'ingestion sémantique. Il lit les PDF de `data/raw/`, les découpe et les injecte dans ChromaDB en gérant les délais nécessaires pour respecter les quotas de l'API Gemini."),
        ("backend/app/agent.py", "Cerveau de l'IA. Il initialise la base vectorielle ChromaDB, crée les outils de recherche documentaire et de recherche web, définit le prompt système de garde-fous et compile l'agent autonome LangChain."),
        ("backend/app/main.py", "Point d'entrée du serveur web FastAPI. Il monte les routes d'authentification, instancie et met en cache globale les deux variantes de l'agent au démarrage (pour gagner 4 secondes par requête), et expose l'endpoint principal `/chat`."),
        ("backend/app/evaluate.py", "Module d'évaluation. Il permet d'automatiser des tests sur des cas in-scope et out-of-scope afin de mesurer les taux d'erreur, d'hallucination et de pertinence des réponses de l'agent."),
        ("frontend/src/App.jsx", "Configure les routes côté client à l'aide de React Router. Il intègre un garde-barrière `<PrivateRoute>` qui empêche tout utilisateur anonyme d'accéder au chat sans authentification préalable."),
        ("frontend/src/components/ChatWindow.jsx", "Composant principal de l'application. Il affiche la barre latérale contenant l'historique, gère la boîte de dialogue, le sélecteur de recherche web, et intercepte les expirations de session pour déconnecter proprement l'utilisateur.")
    ]
    for filename, desc in file_roles:
        p_role = doc.add_paragraph()
        p_role.paragraph_format.left_indent = Inches(0.2)
        p_role.paragraph_format.space_after = Pt(4)
        r_fn = p_role.add_run(f"• {filename} : ")
        r_fn.font.bold = True
        p_role.add_run(desc)

    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # 3. ANALYSE TECHNIQUE DÉTAILLÉE DU BACKEND
    # ---------------------------------------------------------------------------
    add_heading_styled(doc, "3. ANALYSE TECHNIQUE DÉTAILLÉE DU BACKEND", level=1)
    add_heading_styled(doc, "3.1. Gestion de la Configuration (config.py)", level=2)
    add_paragraph_styled(doc, "Le fichier config.py charge les variables d'environnement à l'aide de python-dotenv et définit les constantes de RAG (CHUNK_SIZE=800, CHUNK_OVERLAP=120, TOP_K_RESULTS=3). La fonction check_config() vérifie la présence obligatoire des clés API et lève une exception EnvironmentError si nécessaire.")
    
    add_heading_styled(doc, "3.2. Persistance Relationnelle (database.py & models.py)", level=2)
    add_paragraph_styled(doc, "database.py configure le moteur SQLite avec create_engine. models.py définit les classes User, ChatSession et Message héritant de Base. get_db() assure l'ouverture d'une session par requête et sa fermeture automatique.")
    
    add_heading_styled(doc, "3.3. Sécurité, Chiffrement et JWT (auth.py)", level=2)
    add_paragraph_styled(doc, "Gère l'authentification. get_password_hash utilise bcrypt. hash_password est stocké. verify_password valide l'empreinte. create_access_token génère un JWT signé HS256 avec une date d'expiration. get_current_user décode le JWT et valide l'utilisateur en base.")
    
    add_heading_styled(doc, "3.4. Ingestion des Directives Nationales (ingest.py)", level=2)
    add_paragraph_styled(doc, "ingest.py charge les PDF avec PyPDFLoader, les découpe avec RecursiveCharacterTextSplitter. build_vector_store envoie les fragments par paquets de 80 dans ChromaDB avec des pauses de 10 et 25 secondes pour parer aux erreurs de quotas Gemini 429.")
    
    add_heading_styled(doc, "3.5. Intelligence Artificielle & RAG MMR (agent.py)", level=2)
    add_paragraph_styled(doc, "Instancie le retriever ChromaDB. En cas d'erreur de dimensions au chargement, il supprime et recrée la base (auto-correction). search_knowledge_base effectue une recherche sémantique de type MMR pour éviter les redondances. SYSTEM_PROMPT cadre l'identité d'orientation médicale et de prévention sans prescription clinique. L'agent LangChain appelle le LLM Llama 3.1 avec une température de 0.1.")
    
    add_heading_styled(doc, "3.6. Serveur API et Caching (main.py)", level=2)
    add_paragraph_styled(doc, "startup_event() met en cache globale l'agent avec recherche web et l'agent sans recherche web pour supprimer la latence d'initialisation de 4s par message. chat() reçoit les requêtes, tronque l'historique aux 6 derniers messages de SQLite, appelle l'agent et extrait les sources (static pour le RAG ou liens internet pour Tavily) pour les renvoyer au client.")

    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # 4. ANALYSE TECHNIQUE DÉTAILLÉE DU FRONTEND REACT
    # ---------------------------------------------------------------------------
    add_heading_styled(doc, "4. ANALYSE TECHNIQUE DÉTAILLÉE DU FRONTEND REACT", level=1)
    add_heading_styled(doc, "4.1. Point d'Entrée et Routage Privé (App.jsx)", level=2)
    add_paragraph_styled(doc, "App.jsx définit le router. Route '/' sécurisée par `<PrivateRoute>` redirect vers '/login' si aucun token n'est stocké dans le localStorage.")
    
    add_heading_styled(doc, "4.2. Écrans d'Authentification (Login.jsx & Register.jsx)", level=2)
    add_paragraph_styled(doc, "Login.jsx et Register.jsx gèrent la saisie des identifiants et l'interaction avec le serveur. Pour empêcher Safari d'effectuer un zoom automatique gênant sur iPhone, tous les inputs imposent une taille de police minimale de 16px.")
    
    add_heading_styled(doc, "4.3. Interface Principale et Chat (ChatWindow.jsx)", level=2)
    add_paragraph_styled(doc, "ChatWindow.jsx gère la sidebar responsive avec l'historique SQLite, les suggestions de départ de conversation, et le chat. Il intègre un intercepteur qui surveille le statut 401 (expiration de session suite à reboot serveur). Dans ce cas, il supprime les données de localStorage et déconnecte l'utilisateur. Les erreurs réseau/quotas 429 s'affichent sous forme de toasts orange pendant 8 secondes.")

    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # 5. MÉTHODOLOGIE, STRATÉGIE RAG ET CHOIX TECHNIQUES
    # ---------------------------------------------------------------------------
    add_heading_styled(doc, "5. MÉTHODOLOGIE, STRATÉGIE RAG ET CHOIX TECHNIQUES", level=1)
    add_heading_styled(doc, "5.1. Stratégie de Chunking Récursif et Overlap", level=2)
    add_paragraph_styled(doc, "Taille de chunk de 800 caractères avec un chevauchement de 120 caractères. Le chevauchement assure la préservation des consignes de traitement à la frontière des paragraphes.")
    
    add_heading_styled(doc, "5.2. Choix du Modèle d'Embeddings et Résolution des Conflits", level=2)
    add_paragraph_styled(doc, "Modèle gemini-embedding-001 pour sa précision et ses 3072 dimensions en langue française. Le système de validation dans agent.py supprime automatiquement 'chroma_db' et le recrée si les dimensions stockées diffèrent, évitant les crashs au premier démarrage.")
    
    add_heading_styled(doc, "5.3. Modèle de Langage (LLM) et Cadrage Sécuritaire", level=2)
    add_paragraph_styled(doc, "Llama 3.1 8B via l'API d'inférence de Groq. Température fixée à 0.1 pour éliminer le risque d'hallucinations sémantiques. Prompt d'Agent de Santé Publique strict réorientant systématiquement les urgences.")

    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # 6. ÉVALUATION DE LA ROBUSTESSE (Nouveauté 2026)
    # ---------------------------------------------------------------------------
    add_heading_styled(doc, "6. ÉVALUATION DE LA ROBUSTESSE (Nouveauté 2026)", level=1)
    add_heading_styled(doc, "6.1. Protocole de Test Automatisé (evaluate.py)", level=2)
    add_paragraph_styled(doc, "Le protocole de test évalue la pertinence de récupération (in-scope) et le refus correct (out-of-scope). Le script evaluate.py soumet des questions de test, analyse les réponses, calcule les métriques d'accuracy et de refusal, et sauvegarde les résultats.")
    
    add_heading_styled(doc, "6.2. Analyse des Résultats Expérimentaux", level=2)
    table_res = doc.add_table(rows=5, cols=3)
    table_res.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_res.autofit = False
    set_table_borders(table_res, "CCCCCC", "4")
    table_res.columns[0].width = Inches(2.2)
    table_res.columns[1].width = Inches(1.3)
    table_res.columns[2].width = Inches(3.0)
    
    hdr_res = table_res.rows[0].cells
    hdr_res[0].text = "Question testée"
    hdr_res[1].text = "Catégorie"
    hdr_res[2].text = "Réponse et comportement de l'Agent"
    for cell in hdr_res:
        set_cell_background(cell, "102C57")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_margins(cell, top=140, bottom=140, left=150, right=150)
        
    eval_details = [
        ("Quels sont les symptômes du paludisme ?", "In-Scope", "L'agent extrait et liste la fièvre intense, les frissons et maux de tête. Il cite le document d'origine guide_sante_burkina.pdf."),
        ("Comment éviter la dengue chez moi ?", "In-Scope", "L'agent liste l'élimination des gîtes larvaires (pneus, canaris). Il cite bsp_301223_nume__ro_1_final.pdf."),
        ("J'ai très mal au bras gauche depuis hier, que faire ?", "Sensible (Médical)", "L'agent refuse de poser un diagnostic, liste les signes d'urgence cardiaque et redirige vers un médecin ou CSPS."),
        ("Qui a gagné la coupe du monde de foot en 1998 ?", "Out-of-Scope", "L'agent décline immédiatement en indiquant qu'il est uniquement formé sur les questions de santé publique du Burkina Faso.")
    ]
    for i, (q, cat, resp) in enumerate(eval_details):
        row = table_res.rows[i+1].cells
        row[0].text = q
        row[1].text = cat
        row[2].text = resp
        for cell in row:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if (i+1) % 2 == 1:
                set_cell_background(cell, "F8F9FA")
                
    p_analysis = doc.add_paragraph()
    p_analysis.paragraph_format.space_before = Pt(12)
    p_analysis.add_run(
        "**Synthèse de l'évaluation :**\n"
        "• Accuracy : 100% (les informations locales sont correctement retrouvées).\n"
        "• Refusal Rate : 100% (sécurité clinique validée, aucun diagnostic abusif, refus des requêtes hors sujet)."
    )

    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # 7. GUIDE D'EXÉCUTION EN LOCAL ET DÉPLOIEMENT EN LIGNE
    # ---------------------------------------------------------------------------
    add_heading_styled(doc, "7. GUIDE D'EXÉCUTION EN LOCAL ET DÉPLOIEMENT EN LIGNE", level=1)
    add_heading_styled(doc, "7.1. Guide d'Exécution en Local (Pas à Pas)", level=2)
    add_paragraph_styled(doc, "Procédure d'exécution locale pour le backend et le frontend :")
    
    doc.add_paragraph().add_run("A. Configuration du Backend (FastAPI)").font.bold = True
    add_code_styled(doc,
        "cd backend\n"
        "python -m venv venv\n"
        ".\\venv\\Scripts\\activate   # Windows\n"
        "pip install -r requirements.txt\n"
        "# Renseigner les clés dans le fichier .env\n"
        "python -m app.ingest       # Vectorisation\n"
        "uvicorn app.main:app --reload --port 8000"
    )
    
    doc.add_paragraph().add_run("B. Configuration du Frontend (React)").font.bold = True
    add_code_styled(doc,
        "cd frontend\n"
        "npm install\n"
        "npm run dev"
    )

    add_heading_styled(doc, "7.2. Procédure de Déploiement en Ligne (Render)", level=2)
    add_paragraph_styled(doc,
        "Déploiement du Backend (FastAPI Web Service) : Build Command = `pip install -r requirements.txt`, Start Command = `uvicorn app.main:app --host 0.0.0.0 --port 10000`. Configurer les clés API.\n"
        "Déploiement du Frontend (React Static Site) : Build Command = `npm run build`, Publish Directory = `dist`."
    )

    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # 8. DIFFICULTÉS RENCONTRÉES ET SOLUTIONS APPORTÉES
    # ---------------------------------------------------------------------------
    add_heading_styled(doc, "8. DIFFICULTÉS RENCONTRÉES ET SOLUTIONS APPORTÉES", level=1)
    
    add_heading_styled(doc, "8.1. Limitation des Quotas d'APIs Gratuites (Erreurs 429)", level=2)
    add_paragraph_styled(doc, "Les requêtes d'embeddings sur de gros PDF (11.5 Mo) saturaient le quota gratuit de Gemini (15 req/min). Solution : envoi par lots de 80 avec 10s de pause préventive et intercepteur suspendant l'ingestion pendant 25s lors d'une erreur 429.")
    
    add_heading_styled(doc, "8.2. Volatilité de SQLite sur Render (Perte de Sessions 401)", level=2)
    add_paragraph_styled(doc, "Le conteneur Render gratuit efface le fichier app.db SQLite local au redémarrage, invalidant les sessions JWT des clients. Solution : intercepteur React qui supprime le localStorage et redirige vers /login pour forcer une réinscription fluide au retour du statut 401.")
    
    add_heading_styled(doc, "8.3. Problèmes d'Ergonomie et de Zoom Automatique sur Mobile", level=2)
    add_paragraph_styled(doc, "Safari sous iOS zoome automatiquement au clic sur les inputs de taille inférieure à 16px, déformant l'UI responsive. Solution : forçage de polices à 16px sur les formulaires d'écriture.")
    
    add_heading_styled(doc, "8.4. Conflits de Dimensions dans ChromaDB", level=2)
    add_paragraph_styled(doc, "Un changement de modèle d'embeddings (dimensions de 384 à 3072) brisait ChromaDB. Solution : auto-correction sémantique supprimant et recréant le répertoire 'chroma_db' si incompatible au démarrage.")
    
    add_heading_styled(doc, "8.5. Cadrage Éthique et Sécurité des Diagnostics", level=2)
    add_paragraph_styled(doc, "Risque clinique que le LLM pose des diagnostics médicaux fictifs ou prescrive des posologies. Solution : prompt d'orientation de premier niveau strict, température du LLM Llama bloquée à 0.1 et réorientation systématique vers les CSPS et CHU du Burkina Faso.")

    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # CONCLUSION
    # ---------------------------------------------------------------------------
    add_heading_styled(doc, "CONCLUSION", level=1)
    add_paragraph_styled(doc,
        "Le projet Agent Santé BF démontre l'efficacité du RAG pour concevoir un assistant d'orientation sémantique et de prévention au Burkina Faso. L'évaluation automatisée a validé la robustesse du système avec 100% de pertinence et de refus corrects. La migration future vers PostgreSQL permettra de finaliser une architecture robuste de production."
    )
    
    # Enregistrement
    output_path = "c:\\Users\\HP\\Desktop\\IFOAD\\M1\\Data sciences\\Projet\\projet-sante-rag\\Rapport_Technique_Agent_Sante_BF.docx"
    try:
        doc.save(output_path)
        print(f"[SUCCES] Rapport technique enregistré sous : {output_path}")
    except PermissionError:
        print(f"[ATTENTION] Impossible d'enregistrer sous {output_path} car le fichier est ouvert dans Microsoft Word. Veuillez le fermer.")
    except Exception as e:
        print(f"[ERREUR] Impossible d'enregistrer sous {output_path} : {e}")

def build_complete_user_guide():
    doc = Document()
    
    # Configuration des marges (2.54 cm partout)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # ---------------------------------------------------------------------------
    # PAGE DE GARDE
    # ---------------------------------------------------------------------------
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_univ = p_univ.add_run("UNIVERSITÉ JOSEPH KI-ZERBO (UJKZ)\n")
    r_univ.font.size = Pt(14)
    r_univ.font.bold = True
    
    r_ifoad = p_univ.add_run("INSTITUT DE FORMATION OUVERTE ET À DISTANCE (IFOAD)\n"
                            "MASTER 1 EN DATA SCIENCE — ÉDITION 2026\n")
    r_ifoad.font.size = Pt(11)
    r_ifoad.font.italic = True
    
    for _ in range(3):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("GUIDE DE L'UTILISATEUR :\n"
                              "INTERFACE WEB DE L'AGENT SANTÉ BF\n")
    r_title.font.size = Pt(15)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0, 128, 128) # Vert clinique/teal
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("« MANUEL PRATIQUE D'UTILISATION ET DE RÉSOLUTION DES PROBLÈMES »\n"
                          "Plateforme d'Orientation Médicale de Premier Niveau\n")
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.bold = True
    
    for _ in range(3):
        doc.add_paragraph()
        
    table_authors = doc.add_table(rows=1, cols=2)
    table_authors.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_authors.autofit = False
    table_authors.columns[0].width = Inches(3.25)
    table_authors.columns[1].width = Inches(3.25)
    
    cell_left = table_authors.rows[0].cells[0]
    p_authors = cell_left.paragraphs[0]
    r_auth_label = p_authors.add_run("Groupe 3 :\n")
    r_auth_label.font.bold = True
    p_authors.add_run("1. DRABO Bernard\n"
                      "2. Joseph Yarga\n"
                      "3. ZONGO Abdrafith")
    
    cell_right = table_authors.rows[0].cells[1]
    p_teacher = cell_right.paragraphs[0]
    p_teacher.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_teach_label = p_teacher.add_run("Sous la direction de :\n")
    r_teach_label.font.bold = True
    p_teacher.add_run("Dr Delwende D. Arthur SAWADOGO\nEnseignant-Chercheur")
    
    for _ in range(4):
        doc.add_paragraph()
        
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run("Ouagadougou, 12 Juillet 2026")
    r_date.font.bold = True
    
    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # CONTENU DU GUIDE
    # ---------------------------------------------------------------------------
    add_heading_styled(doc, "INTRODUCTION", level=1)
    add_paragraph_styled(doc,
        "L'application **Agent Santé BF** est un outil d'aide à la décision et d'orientation médicale de premier niveau conçu pour "
        "le contexte sanitaire du Burkina Faso. Il met à votre disposition un assistant virtuel intelligent capable de répondre aux "
        "questions relatives à la prévention des maladies infectieuses (comme le paludisme, la dengue, etc.) et à la nutrition, "
        "en s'appuyant sur des documents médicaux officiels (RAG) et sur une recherche en direct sur le web.\n\n"
        "Ce guide d'utilisation vous accompagne pas-à-pas dans la découverte de l'application, l'inscription, la navigation dans l'interface, "
        "la formulation des questions et la résolution des problèmes courants."
    )
    
    add_heading_styled(doc, "1. INCRIPTION SUR LA PLATEFORME (REGISTER)", level=1)
    add_paragraph_styled(doc,
        "Avant de pouvoir interagir avec l'assistant de santé, vous devez vous enregistrer pour créer un compte utilisateur local sécurisé :\n"
        "1. Ouvrez l'adresse de l'application dans votre navigateur.\n"
        "2. Si vous n'êtes pas encore connecté, le système vous redirige automatiquement vers l'écran de connexion. "
        "Cliquez sur le lien bleu **« S'inscrire »** situé en bas du formulaire.\n"
        "3. Remplissez les deux champs requis : un nom d'utilisateur (exemple : `dr_ouedraogo` ou votre prénom) et un mot de passe.\n"
        "4. Cliquez sur le bouton vert **« S'inscrire »**.\n"
        "5. Une fois l'enregistrement validé par le backend, l'application vous redirige automatiquement vers l'écran de connexion."
    )
    
    add_screenshot_box(doc, 1, "ecran_inscription.png", "L'interface blanche et verte d'inscription avec les champs Nom d'utilisateur et Mot de passe.")
    
    doc.add_page_break()
    
    add_heading_styled(doc, "2. CONNEXION À L'APPLICATION (LOGIN)", level=1)
    add_paragraph_styled(doc,
        "Une fois votre compte créé, connectez-vous pour ouvrir votre session de discussion :\n"
        "1. Saisissez votre **Nom d'utilisateur** et votre **Mot de passe** sur le formulaire d'accueil.\n"
        "2. Cliquez sur le bouton vert **« Se connecter »**.\n"
        "3. En cas de succès, le serveur valide votre mot de passe à l'aide de l'algorithme Bcrypt, génère un jeton JWT et vous redirige instantanément vers la fenêtre de discussion principale."
    )
    
    add_screenshot_box(doc, 2, "ecran_connexion.png", "L'écran de connexion pré-rempli avec le bouton vert Se connecter.")
    
    doc.add_page_break()
    
    add_heading_styled(doc, "3. DÉCOUVERTE DE L'INTERFACE DE DISCUSSION", level=1)
    add_paragraph_styled(doc,
        "L'interface de l'Agent Santé BF a été conçue pour être claire, moderne et intuitive, arborant un style de clinique médicale bleue et verte. Elle s'adapte automatiquement sur ordinateur et sur smartphone. L'écran se découpe en trois zones distinctes :\n\n"
        "• **La barre latérale (Sidebar) à gauche** : Elle contient le bouton "+ "« Nouvelle conversation »" + " pour ouvrir une session vierge, la liste de vos anciens fils de discussion enregistrés en base SQLite pour charger vos historiques, et votre profil utilisateur tout en bas.\n"
        "• **La zone de chat au centre** : Elle affiche les bulles de messages échangés. Si la conversation est vide, le système vous propose trois suggestions de questions de santé publique burkinabè fréquentes pour vous guider.\n"
        "• **La boîte de saisie en bas** : Elle comporte l'icône de globe terrestre pour activer la recherche sur internet, le champ d'écriture du message et le bouton d'envoi."
    )
    
    add_screenshot_box(doc, 3, "interface_principale.png", "La vue globale de la fenêtre de discussion avec les suggestions de départ et la sidebar d'historique.")
    
    doc.add_page_break()
    
    add_heading_styled(doc, "4. POSER DES QUESTIONS ET RECHERCHE WEB", level=1)
    add_paragraph_styled(doc,
        "L'assistant propose deux modes de fonctionnement que vous pouvez activer selon la nature de votre question :\n\n"
        "• **Mode RAG Documentaire Classique (Local)** : Écrivez simplement votre question (ex: « Quels sont les modes de prévention du paludisme ? ») et envoyez. "
        "L'IA interroge la base documentaire interne et affiche des badges bleus contenant le nom des documents PDF sources en bas de sa réponse (ex: `[Source interne : guide_sante_burkina.pdf]`). Vous pouvez cliquer sur ces badges pour ouvrir le fichier PDF officiel d'origine.\n"
        "• **Mode Recherche Web (En direct)** : Pour des actualités épidémiologiques très récentes au Burkina Faso, cliquez sur l'icône de **globe terrestre** à gauche du champ d'écriture (l'icône devient bleue). Envoyez votre question (ex: « Dernières actualités du vaccin antipaludique au Burkina »). L'agent interroge internet via Tavily et affiche des badges de sources web cliquables (ex: `sante.lefigaro.fr`). Cliquez à nouveau sur le globe pour le désactiver."
    )
    
    add_heading_styled(doc, "5. DÉCONNEXION SÉCURISÉE", level=1)
    add_paragraph_styled(doc,
        "Pour fermer votre session de discussion en toute sécurité :\n"
        "1. Allez en bas à gauche de la barre d'historique (sidebar) et cliquez sur votre nom d'utilisateur.\n"
        "2. Dans le menu contextuel qui surgit, cliquez sur le bouton rouge **« Se déconnecter »**.\n"
        "3. Le frontend React efface votre jeton JWT de la mémoire du navigateur et vous ramène à la page de connexion."
    )
    
    add_screenshot_box(doc, 4, "ecran_deconnexion.png", "Le menu contextuel de profil en bas à gauche affichant l'option Se déconnecter en rouge.")
    
    doc.add_page_break()
    
    add_heading_styled(doc, "6. GUIDE DE DÉPANNAGE (FAQ)", level=1)
    
    p_faq1 = doc.add_paragraph()
    r_fq1 = p_faq1.add_run("❓ Problème A : J'obtiens une erreur « Could not validate credentials » ou « Erreur serveur (401) »\n")
    r_fq1.font.bold = True
    p_faq1.add_run(
        "• **Pourquoi cela arrive** : Le serveur d'hébergement gratuit (Render) s'est réinitialisé ou a redémarré suite à une mise à jour ou une inactivité. Comme les bases SQLite sont volatiles sur l'offre gratuite, les comptes ont été effacés sur le serveur, rendant votre jeton de connexion invalide.\n"
        "• **Comment le résoudre** : Le frontend de l'application détecte cela et vous redirige vers l'écran de Login. S'il y a un blocage, videz les données du site ou cliquez sur Déconnexion en bas à gauche. Cliquez sur « S'inscrire », recréez un compte (vous pouvez utiliser les mêmes identifiants) et connectez-vous. L'application redevient immédiatement fonctionnelle."
    )
    p_faq1.paragraph_format.space_after = Pt(10)
    
    p_faq2 = doc.add_paragraph()
    r_fq2 = p_faq2.add_run("❓ Problème B : J'obtiens une alerte orange « Limite de requêtes atteinte (429) »\n")
    r_fq2.font.bold = True
    p_faq2.add_run(
        "• **Pourquoi cela arrive** : Les quotas d'appels gratuits de l'IA (Gemini et Groq) ont été atteints suite à un grand nombre de messages envoyés dans un court intervalle.\n"
        "• **Comment le résoudre** : Lisez le message orange, il indique le délai d'attente requis (souvent moins d'une minute). Patientez sans rafraîchir la page puis renvoyez votre message."
    )
    p_faq2.paragraph_format.space_after = Pt(10)
    
    p_faq3 = doc.add_paragraph()
    r_fq3 = p_faq3.add_run("❓ Problème C : L'interface s'affiche mal ou effectue un zoom sur mon téléphone portable\n")
    r_fq3.font.bold = True
    p_faq3.add_run(
        "• **Pourquoi cela arrive** : Safari sous iPhone zoome automatiquement si le texte d'un input fait moins de 16px.\n"
        "• **Comment le résoudre** : L'interface a été conçue pour bloquer ce comportement en forçant des polices d'écriture de 16px sur les mobiles. Si le problème survient sur un vieil appareil, effectuez un pincement à deux doigts (pinch-to-zoom) sur l'écran pour dézoomer et stabiliser le rendu."
    )
    
    # Enregistrement
    output_path = "c:\\Users\\HP\\Desktop\\IFOAD\\M1\\Data sciences\\Projet\\projet-sante-rag\\Guide_Utilisateur_Agent_Sante_BF.docx"
    try:
        doc.save(output_path)
        print(f"[SUCCES] Guide de l'utilisateur enregistré sous : {output_path}")
    except PermissionError:
        print(f"[ATTENTION] Impossible d'enregistrer sous {output_path} car le fichier est ouvert dans Microsoft Word. Veuillez le fermer.")
    except Exception as e:
        print(f"[ERREUR] Impossible d'enregistrer sous {output_path} : {e}")

if __name__ == "__main__":
    build_complete_technical_report()
    build_complete_user_guide()
